# Copyright (c) 2025 Dedalus Labs, Inc. and its contributors
# SPDX-License-Identifier: MIT

"""PDF generation tools for MCP server.

Generates real PDF files from markdown content, demonstrating
MCP's ability to provide capabilities models can't do natively.
"""

import hashlib
import httpx
import os
import re
import secrets
import tempfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from pydantic import BaseModel
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    ListFlowable,
    ListItem,
)

from dedalus_mcp import tool


# --- Configuration -----------------------------------------------------------

# Directory to store generated files
FILES_DIR = Path(tempfile.gettempdir()) / "pdf-generator-mcp"
FILES_DIR.mkdir(exist_ok=True)

# Base URL for file downloads (file server runs on separate port)
BASE_URL = os.getenv("BASE_URL", "http://127.0.0.1:8081")


# --- Response Models ---------------------------------------------------------


class PdfResult(BaseModel):
    """Result of PDF generation."""

    success: bool
    pdf_id: str | None = None
    filename: str | None = None
    size_bytes: int | None = None
    download_url: str | None = None
    error: str | None = None


class DocxResult(BaseModel):
    """Result of DOCX generation."""

    success: bool
    docx_id: str | None = None
    filename: str | None = None
    size_bytes: int | None = None
    download_url: str | None = None
    error: str | None = None


# --- Styles ------------------------------------------------------------------

def get_styles(style_name: str = "default"):
    """Get paragraph styles based on style name."""
    styles = getSampleStyleSheet()

    if style_name == "modern":
        title_color = colors.HexColor("#111827")
        heading_color = colors.HexColor("#374151")
        text_color = colors.HexColor("#1f2937")
        accent_color = colors.HexColor("#3b82f6")
    elif style_name == "minimal":
        title_color = colors.black
        heading_color = colors.black
        text_color = colors.HexColor("#222222")
        accent_color = colors.HexColor("#666666")
    else:  # default
        title_color = colors.HexColor("#1a1a1a")
        heading_color = colors.HexColor("#1a1a1a")
        text_color = colors.HexColor("#333333")
        accent_color = colors.HexColor("#3b82f6")

    custom_styles = {
        "Title": ParagraphStyle(
            "CustomTitle",
            parent=styles["Title"],
            fontSize=24,
            textColor=title_color,
            spaceAfter=20,
            borderPadding=(0, 0, 8, 0),
            borderWidth=0,
            borderColor=accent_color,
        ),
        "Heading1": ParagraphStyle(
            "CustomH1",
            parent=styles["Heading1"],
            fontSize=18,
            textColor=heading_color,
            spaceBefore=20,
            spaceAfter=12,
        ),
        "Heading2": ParagraphStyle(
            "CustomH2",
            parent=styles["Heading2"],
            fontSize=14,
            textColor=heading_color,
            spaceBefore=16,
            spaceAfter=8,
        ),
        "Heading3": ParagraphStyle(
            "CustomH3",
            parent=styles["Heading3"],
            fontSize=12,
            textColor=heading_color,
            spaceBefore=12,
            spaceAfter=6,
        ),
        "BodyText": ParagraphStyle(
            "CustomBody",
            parent=styles["BodyText"],
            fontSize=11,
            textColor=text_color,
            spaceBefore=6,
            spaceAfter=6,
            leading=16,
        ),
        "BulletText": ParagraphStyle(
            "CustomBullet",
            parent=styles["BodyText"],
            fontSize=11,
            textColor=text_color,
            leftIndent=20,
            spaceBefore=4,
            spaceAfter=4,
        ),
    }

    return custom_styles, accent_color


# --- Helper Functions --------------------------------------------------------


def _generate_id(content: str) -> str:
    """Generate a cryptographically secure unique ID.

    Uses secrets module for secure random generation plus content hash
    to create an unguessable 32-character token.
    """
    # 16 bytes of cryptographically secure randomness = 32 hex chars
    random_token = secrets.token_hex(16)
    # Add content hash for uniqueness
    content_hash = hashlib.sha256(f"{content}{datetime.now().isoformat()}".encode()).hexdigest()[:8]
    return f"{random_token}{content_hash}"


def _slugify(title: str) -> str:
    """Convert title to a safe filename."""
    slug = title.lower()
    slug = "".join(c if c.isalnum() or c == " " else "" for c in slug)
    slug = "-".join(slug.split())
    return slug[:50] or "document"


def _parse_markdown_to_flowables(markdown_content: str, styles: dict, accent_color):
    """Convert markdown to reportlab flowables."""
    flowables = []
    lines = markdown_content.strip().split("\n")

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            flowables.append(Spacer(1, 6))
            i += 1
            continue

        # Headers
        if line.startswith("### "):
            flowables.append(Paragraph(line[4:], styles["Heading3"]))
        elif line.startswith("## "):
            flowables.append(Paragraph(line[3:], styles["Heading2"]))
        elif line.startswith("# "):
            flowables.append(Paragraph(line[2:], styles["Heading1"]))

        # Bold text (convert **text** to <b>text</b>)
        elif "**" in line:
            formatted = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
            flowables.append(Paragraph(formatted, styles["BodyText"]))

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            bullet_items = []
            while i < len(lines) and (lines[i].strip().startswith("- ") or lines[i].strip().startswith("* ")):
                item_text = lines[i].strip()[2:]
                # Handle bold in list items
                item_text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", item_text)
                bullet_items.append(ListItem(Paragraph(item_text, styles["BulletText"])))
                i += 1
            flowables.append(ListFlowable(bullet_items, bulletType="bullet", leftIndent=20))
            continue

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            num_items = []
            while i < len(lines) and re.match(r"^\d+\. ", lines[i].strip()):
                item_text = re.sub(r"^\d+\. ", "", lines[i].strip())
                item_text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", item_text)
                num_items.append(ListItem(Paragraph(item_text, styles["BulletText"])))
                i += 1
            flowables.append(ListFlowable(num_items, bulletType="1", leftIndent=20))
            continue

        # Table (simple markdown table support)
        elif line.startswith("|"):
            table_data = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_line = lines[i].strip()
                # Skip separator line
                if re.match(r"^\|[-:\s|]+\|$", row_line):
                    i += 1
                    continue
                # Parse cells
                cells = [cell.strip() for cell in row_line.split("|")[1:-1]]
                table_data.append(cells)
                i += 1

            if table_data:
                table = Table(table_data, repeatRows=1)
                table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f8f9fa")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#374151")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 10),
                    ("TOPPADDING", (0, 0), (-1, 0), 10),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 1), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 1), (-1, -1), 8),
                ]))
                flowables.append(Spacer(1, 10))
                flowables.append(table)
                flowables.append(Spacer(1, 10))
            continue

        # Regular paragraph
        else:
            formatted = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
            flowables.append(Paragraph(formatted, styles["BodyText"]))

        i += 1

    return flowables


# --- PDF Tools ---------------------------------------------------------------


@tool(
    description=(
        "Generate a PDF document from markdown content. "
        "Returns a download URL for the generated PDF file. "
        "Supports three styles: 'default' (professional blue accents), "
        "'modern' (clean contemporary design), 'minimal' (elegant serif)."
    )
)
async def render_pdf(
    title: str,
    markdown: str,
    style: str = "default",
) -> PdfResult:
    """Generate a PDF from markdown content.

    Args:
        title: Document title (appears as heading and in filename).
        markdown: Markdown content for the document body.
        style: Visual style - 'default', 'modern', or 'minimal'.

    Returns:
        PdfResult with pdf_id, filename, size_bytes, and download_url.
    """
    try:
        # Validate style
        if style not in ["default", "modern", "minimal"]:
            style = "default"

        # Generate cryptographically secure ID and filename
        # Using only the secure token (no predictable slug) for security
        pdf_id = _generate_id(f"{title}{markdown}")
        filename = f"{pdf_id}.pdf"
        filepath = FILES_DIR / filename

        # Get styles
        custom_styles, accent_color = get_styles(style)

        # Create PDF document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        # Build content
        flowables = []

        # Add title
        flowables.append(Paragraph(title, custom_styles["Title"]))
        flowables.append(Spacer(1, 12))

        # Parse markdown content
        content_flowables = _parse_markdown_to_flowables(markdown, custom_styles, accent_color)
        flowables.extend(content_flowables)

        # Build PDF
        doc.build(flowables)

        # Get file size
        size_bytes = filepath.stat().st_size

        # Upload to tmpfiles.org for temporary public URL
        async with httpx.AsyncClient() as client:
            with open(filepath, "rb") as f:
                response = await client.post(
                    "https://tmpfiles.org/api/v1/upload",
                    files={"file": (filename, f, "application/pdf")},
                    timeout=30.0,
                )
                response.raise_for_status()
                result = response.json()
                # Convert page URL to direct download URL (add /dl/ to path)
                url = result["data"]["url"]
                download_url = url.replace("tmpfiles.org/", "tmpfiles.org/dl/")

        return PdfResult(
            success=True,
            pdf_id=pdf_id,
            filename=filename,
            size_bytes=size_bytes,
            download_url=download_url,
        )

    except Exception as e:
        return PdfResult(success=False, error=f"PDF generation failed: {e!s}")


@tool(
    description=(
        "Generate a DOCX (Word) document from markdown content. "
        "Returns a download URL for the generated DOCX file."
    )
)
async def render_docx(
    title: str,
    markdown: str,
) -> DocxResult:
    """Generate a DOCX from markdown content.

    Args:
        title: Document title.
        markdown: Markdown content for the document body.

    Returns:
        DocxResult with docx_id, filename, size_bytes, and download_url.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Generate cryptographically secure ID and filename
        docx_id = _generate_id(f"{title}{markdown}")
        filename = f"{docx_id}.docx"
        filepath = FILES_DIR / filename

        # Create document
        doc = Document()

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Parse markdown and add content
        lines = markdown.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # Headers
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            # List items
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
            # Regular paragraph
            else:
                doc.add_paragraph(line)

            i += 1

        # Save document
        doc.save(filepath)

        # Get file size
        size_bytes = filepath.stat().st_size

        # Upload to tmpfiles.org for temporary public URL
        async with httpx.AsyncClient() as client:
            with open(filepath, "rb") as f:
                response = await client.post(
                    "https://tmpfiles.org/api/v1/upload",
                    files={"file": (filename, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                    timeout=30.0,
                )
                response.raise_for_status()
                result = response.json()
                # Convert page URL to direct download URL (add /dl/ to path)
                url = result["data"]["url"]
                download_url = url.replace("tmpfiles.org/", "tmpfiles.org/dl/")

        return DocxResult(
            success=True,
            docx_id=docx_id,
            filename=filename,
            size_bytes=size_bytes,
            download_url=download_url,
        )

    except ImportError:
        return DocxResult(
            success=False,
            error="python-docx not installed. Install with: pip install python-docx",
        )
    except Exception as e:
        return DocxResult(success=False, error=f"DOCX generation failed: {e!s}")


# --- Export ------------------------------------------------------------------

pdf_tools = [
    render_pdf,
    render_docx,
]
